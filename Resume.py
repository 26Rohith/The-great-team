#resume
import docx
import os
def run():
    directory = os.getcwd()
    resume=docx.Document()
    ##resume.save("C:\\Users\\Rohit\\Desktop\\resume maker\\resumenew.docx")
    print("Greetings user")
    print("enter your name")
    user_name=input()
    resume.add_heading(f"Resume of {user_name}",0)
    print(f"Hello {user_name} Let us start building a resume to your needs")
    print("Address:")
    print("Do.No:")
    user_doornum=input()
    print("Area/Landmark:")
    user_area=input()
    print("City/Town:")
    user_city=input()
    print("District:")
    user_district=input()
    print("Pincode:")
    user_pincode=input()
    print("State:")
    user_state=input()
    print("Country:")
    user_country=input()
    resume.add_paragraph(f"""Address: {user_doornum}, {user_area}, {user_city},
    {user_district}, {user_pincode},{user_state}, {user_country}""")
    print("Enter Contact number:")
    user_contact=input()
    print("Enter Email:")
    user_email=input()
    resume.add_paragraph(f"Contact number: {user_contact}")
    resume.add_paragraph(f"Email: {user_email}")
    resume.add_heading("Profile:",1)
    print("A few lines about you.\n Suggetion: Write about things which are use full to the Interviewer. Enter '...' to stop entering.")
    user_profile=""
    while(True):
        ip=input()
        user_profile=user_profile+ip
        if "..." in ip:
            break
    resume.add_paragraph(user_profile.capitalize())
    resume.add_heading("Education:")
    print("Enter School name:")
    user_school=input()
    print("Enter Points Obtained in 10th class:")
    user_tenth_points=input()
    print("Passed out year")
    user_tenth_passedout=input()
    resume.add_paragraph(f"SSC in {user_school} with {user_tenth_points} in year {user_tenth_passedout}")
    resume.add_heading("Work experience:",1)
    print("Add a few lines about your work experience:","Enter '...' to stop entering.")
    user_work=""
    while(True):
        ip=input()
        user_work=user_work+ip
        if "..." in ip:
            break
    resume.add_paragraph(user_work)
    resume.add_heading("Skills",1)
    print("Add a few lines about your Skills:","Enter '...' to stop entering.")
    user_skill=""
    while(True):
        ip=input()
        user_skill=user_skill+ip
        if "..." in ip:
            break
    resume.add_paragraph(user_skill)
    print("Loading.....")
    print("Saved")
    resume.save(f"{directory}\\{user_name}resume.docx")
